from pathlib import Path

from pypdf import PdfReader
from docx import Document

class ResumeReader:


# it will be useful when the file is on the server 
    @staticmethod
    def read_resume(file_path: str = "data/resume.txt") -> str:
        path = Path(file_path)

       
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {path.resolve()}")


        ext = path.suffix.lower()

        if ext ==".txt":
            return ResumeReader.read_txt(path)
        elif ext ==".pdf":
            return ResumeReader.read_pdf(path)
        elif ext in [".doc",".docx"]:
            return ResumeReader.read_docx(path)
        else : 
            raise ValueError(f"Unsupported file '{ext}'. Pls support file in .txt , .pdf , .docx .Or write directly in the text box ")

    @staticmethod
    def read_resume_stream(file_storage)-> str:
        filename = file_storage.filename.lower()

        if filename.endswith(".txt"):
            return file_storage.read().decode("utf-8",errors="ignore").strip()
        elif filename.endswith(".pdf"):
            reader=PdfReader(file_storage)
            text = "\n".join([ page.extract_text()  for page in reader.pages if page.extract_text()])
            return text.strip()

        elif filename.endswith(".docx"):
            doc = Document(file_storage)
            text = "\n".join([para.text for para  in doc.paragraphs if para.text])
            return text.strip()
          
        else:
            raise ValueError("❌ Unsupported file format. Please upload a .txt, .pdf, or .docx file.")


    
    @staticmethod
    def read_txt(path: Path) -> str:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()

    @staticmethod
    def read_pdf(path: Path) -> str:
        reader = PdfReader(str(path))
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        return text.strip()

    @staticmethod
    def read_docx(path: Path) -> str:
        doc = Document(str(path))
        text = "\n".join([para.text for para in doc.paragraphs if para.text])
        return text.strip()


     